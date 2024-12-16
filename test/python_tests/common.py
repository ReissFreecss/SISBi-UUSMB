import docx
from playwright.sync_api import Playwright, sync_playwright, expect
import dotenv, os, pathlib
dotenv.load_dotenv()
raiz = pathlib.Path(os.path.join(os.getcwd(), __file__)).parent.parent.parent/"doc"/"PlantillasReportes"

def login(page):
    page.goto("http://local.uusmb.unam.mx:8080/SISBI/")
    page.get_by_role("button", name="Iniciar sesión").click()
    page.get_by_label("Usuario o correo electrónico:*").click()
    page.get_by_label("Usuario o correo electrónico:*").fill(os.environ["TEST_USERNAME"])
    page.get_by_label("Contraseña*").click()
    page.get_by_label("Contraseña*").fill(os.environ["TEST_PASSWORD"])
    page.get_by_label("Iniciar sesión").get_by_role("button", name="Iniciar sesión").click()

def escoge_proyecto(page, project_ID):
    page.get_by_placeholder("ID o nombre del proyecto").click()
    page.get_by_placeholder("ID o nombre del proyecto").fill(project_ID)
    page.get_by_placeholder("ID o nombre del proyecto").press("Enter")
    page.get_by_role("gridcell", name=project_ID).click()


class TipoMetodologiaYMuestra:
    def __init__(self, tipo_metodologia, tipo_muestra, zipfile):
        self.tipo_metodologia = tipo_metodologia
        self.tipo_muestra = tipo_muestra
        nombarch = f"{tipo_metodologia}{tipo_muestra}.docx"
        if zipfile is None:
             self.document = docx.Document(raiz/nombarch)
        else:
            with zipfile.open(nombarch) as openfile:
                #print("rcompl", nombarch)
                self.document = docx.Document(openfile)

    def printRuns(self):
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                print("run", run.text)
        for table in self.document.tables:
            print("tbl", table)

    def get_valued_pargagraphs(self):
        for paragraph in self.document.paragraphs:
            if pts := paragraph.text.strip():
                yield pts

    def align_texts(self, other):
        valued_paragraphs_self = set(self.get_valued_pargagraphs())
        valued_paragraphs_other = set(other.get_valued_pargagraphs())
        newself = valued_paragraphs_self - valued_paragraphs_other
        newother = valued_paragraphs_other - valued_paragraphs_self
        common = valued_paragraphs_self & valued_paragraphs_other
        if len(newself) or len(newother):
            print(f"differences: {len(newself)} - {len(newother)} (len(common)={len(common)})")
            print("ns: ", "\n    ".join(newself))
            print("no: ", "\n    ".join(newother))
            #print("cm: ", "\n    ".join(common))
